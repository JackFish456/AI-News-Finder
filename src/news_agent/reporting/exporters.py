from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from news_agent.models.cluster import StoryCluster
from news_agent.models.item import ContentItem
from news_agent.summarization.daily_brief import BriefEntry, DailyBriefReport


def _top_stories_heading(report_cfg: dict[str, Any] | None) -> str:
    n = int((report_cfg or {}).get("top_stories", 7))
    return f"Top {n} AI stories (most important)"


def _word_links_inline(doc: Any, urls: list[str]) -> None:
    """Single paragraph: bold 'Links: ' then hyperlinks (space-separated), no extra indent."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    paragraph = doc.add_paragraph()
    part = paragraph.part
    label_run = paragraph.add_run("Links: ")
    label_run.bold = True

    p_el = paragraph._p
    for idx, url in enumerate(urls):
        if idx > 0:
            gap = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = " "
            gap.append(t)
            p_el.append(gap)

        r_id = part.relate_to(str(url), RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = str(url)
        run.append(text_el)
        hyperlink.append(run)
        p_el.append(hyperlink)


def _word_labeled_line(doc: Any, label: str, body: str) -> None:
    """One paragraph: bold `label` (include trailing space if needed) then normal `body`."""
    p = doc.add_paragraph()
    label_run = p.add_run(label)
    label_run.bold = True
    p.add_run(body)


def _add_word_section(doc: Any, title: str, entries: list[BriefEntry]) -> None:
    doc.add_heading(title, level=1)
    if not entries:
        doc.add_paragraph("No items.")
        return
    for i, e in enumerate(entries, start=1):
        doc.add_heading(f"{i}. {e.headline}", level=2)
        _word_labeled_line(doc, "Why it matters: ", e.why_it_matters)
        doc.add_paragraph(e.summary)
        if e.supporting_links:
            _word_links_inline(doc, [str(u) for u in e.supporting_links])
        _word_labeled_line(doc, "Credibility: ", e.credibility_note)
        _word_labeled_line(doc, "Estimated impact: ", e.estimated_impact)


def export_docx(
    brief: DailyBriefReport,
    path: Path,
    *,
    report_cfg: dict[str, Any] | None = None,
) -> None:
    """Write a Word document for easy reading in Microsoft Word."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("AI News Brief", 0)

    rc = report_cfg or {}
    _add_word_section(doc, _top_stories_heading(rc), brief.top_stories)

    doc.save(str(path))


def _section_md(title: str, entries: list[BriefEntry]) -> str:
    lines = [f"## {title}", ""]
    if not entries:
        lines.append("_No items._")
        lines.append("")
        return "\n".join(lines)
    for i, e in enumerate(entries, start=1):
        lines.append(f"### {i}. {e.headline}")
        lines.append("")
        lines.append(f"**Why it matters:** {e.why_it_matters}")
        lines.append("")
        lines.append(e.summary)
        lines.append("")
        if e.supporting_links:
            link_bits = " ".join(f"<{u}>" for u in e.supporting_links)
            lines.append(f"**Links:** {link_bits}")
            lines.append("")
        lines.append(f"**Credibility:** {e.credibility_note}")
        lines.append("")
        lines.append(f"**Estimated impact:** {e.estimated_impact}")
        lines.append("")
    return "\n".join(lines)


def export_markdown(
    brief: DailyBriefReport,
    path: Path,
    *,
    clusters: list[StoryCluster] | None = None,
    audit_items: list[ContentItem] | None = None,
    report_cfg: dict[str, Any] | None = None,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# AI News Brief",
        "",
        _section_md(_top_stories_heading(report_cfg), brief.top_stories),
    ]
    if clusters:
        lines.append("## Cluster map (debug)")
        lines.append("")
        for c in clusters:
            lines.append(f"- **{c.cluster_id}** canonical `{c.canonical_item_id}` — {len(c.member_item_ids)} items")
        lines.append("")
    if audit_items:
        lines.append("## Pipeline audit (compact)")
        lines.append("")
        for it in audit_items[:200]:
            dec = it.pipeline_decision or ""
            sc = it.scores.final_score if it.scores else None
            lines.append(
                f"- `{it.id}` **{dec}** score={sc} — {it.title or it.source_id} — {it.url}"
            )
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def default_output_stem(now: datetime | None = None) -> str:
    now = now or datetime.utcnow()
    return f"brief_{now:%Y%m%d_%H%M}"
